import json
import unittest
from unittest.mock import patch

from app import app, convert_punctuation, layout_settings, process_paragraph
from desktop import DesktopApi, main
from docx import Document


class AppTestCase(unittest.TestCase):
    def setUp(self):
        self.client = app.test_client()

    def test_index_is_available(self):
        response = self.client.get('/')

        self.assertEqual(response.status_code, 200)
        self.assertIn('排版工厂', response.get_data(as_text=True))
        self.assertIn('AutoWord Formatter', response.get_data(as_text=True))

    def test_process_rejects_non_docx_uploads(self):
        response = self.client.post(
            '/process_stream',
            data={'file': (__import__('io').BytesIO(b'not a document'), 'notes.txt')},
            content_type='multipart/form-data',
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn('无有效docx文件', json.loads(response.get_data())['msg'])

    def test_process_stream_saves_upload_before_streaming(self):
        document = Document()
        document.add_paragraph('测试，AutoWord！')
        source = __import__('io').BytesIO()
        document.save(source)
        source.seek(0)

        response = self.client.post(
            '/process_stream',
            data={
                'file': (source, '测试.docx'),
                'language': 'zh',
                'punctuation': 'halfwidth',
            },
            content_type='multipart/form-data',
        )

        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        self.assertNotIn('"status": "error"', body)
        self.assertIn('"status": "done"', body)
        self.assertIn('"val": 1.0', body)

    def test_frontend_resets_button_after_completed_download(self):
        html = self.client.get('/').get_data(as_text=True)

        self.assertIn("startBtn.disabled = false", html)
        self.assertIn("document.getElementById('fileInput').value = ''", html)
        self.assertIn("startBtn.style.backgroundColor = ''", html)
        self.assertIn("window.pywebview.api.export_job", html)
        self.assertIn('id="exportDirectory"', html)

    def test_custom_layout_settings_are_bounded(self):
        settings = layout_settings({
            'paper_size': 'B5', 'custom_margins': 'on', 'top_margin': '99',
            'font_size': '2', 'line_spacing': '1.5', 'punctuation': 'fullwidth',
        })

        self.assertEqual(settings['top'], 8)
        self.assertEqual(settings['font_size'], 6)
        self.assertEqual(settings['line_spacing'], 1.5)
        self.assertEqual(settings['punctuation'], 'fullwidth')

    def test_margin_presets_override_paper_margins(self):
        settings = layout_settings({'margin_preset': 'symmetric'})

        self.assertEqual((settings['top'], settings['bottom'], settings['left'], settings['right']), (0.7, 0.7, 1.5, 0.7))
        self.assertTrue(settings['mirror'])

    def test_language_is_preserved_in_layout_settings(self):
        self.assertEqual(layout_settings({'language': 'en'})['language'], 'en')

    @patch('desktop.webview.start')
    @patch('desktop.webview.create_window')
    def test_desktop_launcher_embeds_backend(self, create_window, start):
        main()

        args, kwargs = create_window.call_args
        self.assertEqual(args[0], 'AutoWord 排版工厂')
        self.assertIs(args[1], app)
        self.assertEqual(kwargs['min_size'], (760, 600))
        self.assertIsInstance(kwargs['js_api'], DesktopApi)
        start.assert_called_once_with(private_mode=True)

    def test_paragraph_formatting_and_punctuation_modes(self):
        document = Document()
        paragraph = document.add_paragraph('Hello, world!')
        settings = layout_settings({'punctuation': 'fullwidth', 'font_size': '12', 'line_spacing': '1.25'})

        process_paragraph(paragraph, settings)

        self.assertEqual(paragraph.text, 'Hello， world！')
        self.assertEqual(paragraph.paragraph_format.line_spacing, 1.25)
        self.assertEqual(paragraph.runs[0].font.size.pt, 12)
        self.assertEqual(convert_punctuation('，。', 'halfwidth'), ',.')
