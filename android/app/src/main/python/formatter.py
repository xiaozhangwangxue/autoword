"""Offline DOCX formatting engine used by the Android application."""
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

HALF = str.maketrans('。，！？：；（）“”‘’【】、', '.,!?:;()""\'\'[],')
FULL = str.maketrans('.,!?:;()"\'[]', '。，！？：；（）“‘【】')


def punctuation(text, mode):
    return text.translate(HALF if mode == 'halfwidth' else FULL) if mode != 'preserve' else text


def page_number(run):
    begin = OxmlElement('w:fldChar'); begin.set(ns.qn('w:fldCharType'), 'begin')
    instruction = OxmlElement('w:instrText'); instruction.set(ns.qn('xml:space'), 'preserve'); instruction.text = 'PAGE'
    end = OxmlElement('w:fldChar'); end.set(ns.qn('w:fldCharType'), 'end')
    run._r.extend((begin, instruction, end))


def format_document(source, destination, raw_options):
    options = json.loads(raw_options)
    doc = Document(source)
    first = punctuation(doc.paragraphs[0].text.strip(), options['punctuation']) if doc.paragraphs else ''
    for paragraph in doc.paragraphs:
        if options['remove_empty'] and not paragraph.text.strip() and not paragraph._element.xpath('.//w:drawing'):
            paragraph._element.getparent().remove(paragraph._element)
            continue
        paragraph.paragraph_format.line_spacing = options['line_spacing']
        paragraph.paragraph_format.space_before = Pt(options['space_before'])
        paragraph.paragraph_format.space_after = Pt(options['space_after'])
        for run in paragraph.runs:
            if run.text:
                run.text = punctuation(run.text, options['punctuation'])
            run.font.size = Pt(options['font_size'])
    for section in doc.sections:
        section.mirror_margins = options.get('mirror', False)
        section.top_margin = Cm(options['top']); section.bottom_margin = Cm(options['bottom'])
        section.left_margin = Cm(options['left']); section.right_margin = Cm(options['right'])
        if options['footer_mode'] != 'none':
            footer = section.footer.paragraphs[0]
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            label = options.get('footer_text') or ('' if options['footer_mode'] == 'page_number' else first)
            if label:
                footer.add_run(label + '  -  ').font.size = Pt(9)
            page_number(footer.add_run())
    doc.save(destination)
    return destination
