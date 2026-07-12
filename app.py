import io
import os
import zipfile
import uuid
import shutil
import json
import time
import threading
from flask import Flask, request, send_file, Response, stream_with_context
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from concurrent.futures import ThreadPoolExecutor

app = Flask(__name__)

# --- 配置区 ---
UPLOAD_FOLDER = '/tmp/mac_pro_uploads'
MAX_UPLOAD_SIZE = int(os.environ.get('MAX_UPLOAD_SIZE', 100 * 1024 * 1024))
app.config['MAX_CONTENT_LENGTH'] = MAX_UPLOAD_SIZE
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# 7天自动清理
FILE_LIFETIME = 7 * 24 * 60 * 60
CLEANUP_INTERVAL = 3600

HALF_WIDTH_TABLE = str.maketrans({
    '。': '.', '，': ',', '！': '!', '？': '?', '：': ':', '；': ';',
    '（': '(', '）': ')', '“': '"', '”': '"', '‘': "'", '’': "'",
    '【': '[', '】': ']', '、': ','
})
FULL_WIDTH_TABLE = str.maketrans({
    '.': '。', ',': '，', '!': '！', '?': '？', ':': '：', ';': '；',
    '(': '（', ')': '）', '"': '“', "'": '‘', '[': '【', ']': '】',
})
PAGE_PRESETS = {
    'A4_STD': (21, 29.7, 0.5, 0.5, 1.5, 0.5, True),
    'A4_EQUAL': (21, 29.7, 0.7, 0.7, 0.7, 0.7, False),
    'B5': (17.6, 25, 0.7, 0.7, 0.7, 0.7, False),
}

# --- 后台清理线程 ---
def auto_cleanup_task():
    while True:
        try:
            now = time.time()
            for filename in os.listdir(UPLOAD_FOLDER):
                file_path = os.path.join(UPLOAD_FOLDER, filename)
                try:
                    if now - os.path.getmtime(file_path) > FILE_LIFETIME:
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            os.unlink(file_path)
                        elif os.path.isdir(file_path):
                            shutil.rmtree(file_path)
                except Exception: pass
        except Exception: pass
        time.sleep(CLEANUP_INTERVAL)

cleanup_thread = threading.Thread(target=auto_cleanup_task, daemon=True)
cleanup_thread.start()

# --- Word 处理逻辑 ---
def create_page_number_field(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def remove_empty_paragraphs(doc):
    count = 0
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[i]
        text = p.text.strip()
        has_image = p._element.xpath('.//w:drawing') or p._element.xpath('.//w:pict')
        if not text and not has_image:
            p._element.getparent().remove(p._element)
            count += 1
    return count

def bounded_number(form, name, default, low, high):
    try:
        return min(max(float(form.get(name, default)), low), high)
    except (TypeError, ValueError):
        return default


def layout_settings(form):
    preset = form.get('paper_size', 'A4_STD')
    if preset not in PAGE_PRESETS:
        preset = 'A4_STD'
    width, height, top, bottom, left, right, mirror = PAGE_PRESETS[preset]
    margin_preset = form.get('margin_preset', 'paper')
    if margin_preset == 'all_05':
        top = bottom = left = right = 0.5
        mirror = False
    elif margin_preset == 'all_07':
        top = bottom = left = right = 0.7
        mirror = False
    elif margin_preset == 'symmetric':
        top = bottom = right = 0.7
        left = 1.5
        mirror = True
    elif form.get('custom_margins') == 'on':
        top = bounded_number(form, 'top_margin', top, 0.1, 8)
        bottom = bounded_number(form, 'bottom_margin', bottom, 0.1, 8)
        left = bounded_number(form, 'left_margin', left, 0.1, 8)
        right = bounded_number(form, 'right_margin', right, 0.1, 8)
        mirror = False
    return {
        'width': width, 'height': height, 'top': top, 'bottom': bottom,
        'left': left, 'right': right, 'mirror': mirror,
        'font_size': bounded_number(form, 'font_size', 10.5, 6, 72),
        'line_spacing': bounded_number(form, 'line_spacing', 1, 0.5, 5),
        'space_before': bounded_number(form, 'space_before', 0, 0, 72),
        'space_after': bounded_number(form, 'space_after', 0, 0, 72),
        'punctuation': form.get('punctuation', 'halfwidth'),
        'remove_empty': form.get('remove_empty') == 'on',
        'footer_mode': form.get('footer_mode', 'first_line'),
        'footer_text': form.get('footer_text', '').strip()[:120],
        'language': form.get('language', 'zh'),
    }


def convert_punctuation(text, mode):
    if mode == 'halfwidth':
        return text.translate(HALF_WIDTH_TABLE)
    if mode == 'fullwidth':
        return text.translate(FULL_WIDTH_TABLE)
    return text


def localized(settings, zh, en):
    return en if settings.get('language') == 'en' else zh


def process_paragraph(paragraph, settings):
    paragraph.paragraph_format.line_spacing = settings['line_spacing']
    paragraph.paragraph_format.space_before = Pt(settings['space_before'])
    paragraph.paragraph_format.space_after = Pt(settings['space_after'])
    for run in paragraph.runs:
        if run.element.xpath('.//w:drawing') or run.element.xpath('.//w:pict'):
            continue
        if run.text:
            text = convert_punctuation(run.text, settings['punctuation'])
            while '\n\n' in text: text = text.replace('\n\n', '\n')
            while '\r\n\r\n' in text: text = text.replace('\r\n\r\n', '\r\n')
            run.text = text
        run.font.size = Pt(settings['font_size'])

def process_single_file_task(filepath, settings, job_id, filename):
    try:
        yield json.dumps({"status": "log", "msg": localized(settings, f"正在读取: {filename}...", f"Reading: {filename}...")}) + "\n"
        doc = Document(filepath)

        first_line_text = ""
        try:
            if len(doc.paragraphs) > 0:
                raw = doc.paragraphs[0].text.strip()
                first_line_text = convert_punctuation(raw, settings['punctuation']) if raw else ""
        except: pass

        if settings['remove_empty']:
            yield json.dumps({"status": "log", "msg": localized(settings, "  ↳ 移除空段落...", "  ↳ Removing empty paragraphs...")}) + "\n"
            remove_empty_paragraphs(doc)

        yield json.dumps({"status": "log", "msg": localized(settings, "  ↳ 标准化排版...", "  ↳ Applying text formatting...")}) + "\n"
        for p in doc.paragraphs: process_paragraph(p, settings)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs: process_paragraph(p, settings)

        yield json.dumps({"status": "log", "msg": localized(settings, "  ↳ 应用页面设置...", "  ↳ Applying page settings...")}) + "\n"
        for section in doc.sections:
            section.header_distance = Cm(0)
            section.footer_distance = Cm(0.7)
            section.page_width = Cm(settings['width']); section.page_height = Cm(settings['height'])
            section.mirror_margins = settings['mirror']
            section.top_margin = Cm(settings['top']); section.bottom_margin = Cm(settings['bottom'])
            section.left_margin = Cm(settings['left']); section.right_margin = Cm(settings['right'])

            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            for p in section.header.paragraphs: p.text = ""
            for p in section.footer.paragraphs: p.text = ""
            if settings['footer_mode'] != 'none':
                footer_para = section.footer.paragraphs[0]
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                footer = settings['footer_text'] or first_line_text
                if settings['footer_mode'] == 'page_number':
                    footer = ''
                if footer:
                    run_text = footer_para.add_run(footer + "  -  ")
                    run_text.font.size = Pt(9)
                create_page_number_field(footer_para.add_run())

        yield json.dumps({"status": "log", "msg": localized(settings, "  ↳ 完成！", "  ↳ Done!")}) + "\n"
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)

        name_part, ext_part = os.path.splitext(filename)
        new_name = f"{name_part}_formatted{ext_part}"
        yield {"filename": new_name, "data": out_stream.getvalue()}

    except Exception as e:
        detail = f"{type(e).__name__}: {e}"
        yield json.dumps({
            "status": "error",
            "msg": localized(
                settings,
                f"❌ 失败: {filename}（{detail}）",
                f"❌ Failed: {filename} ({detail})",
            ),
        }) + "\n"
        yield None

@app.route('/', methods=['GET'])
def index():
    return '''
    <!doctype html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
        <title>Mac Pro 排版工厂</title>
        <style>
            /* 全局重置，适配移动端 */
            * { box-sizing: border-box; -webkit-tap-highlight-color: transparent; }
            body {
                background:
                    radial-gradient(circle at 15% 10%, rgba(90, 175, 255, .42), transparent 34%),
                    radial-gradient(circle at 90% 70%, rgba(126, 94, 255, .3), transparent 38%),
                    linear-gradient(145deg, #dbefff 0%, #f5f1ff 48%, #d9edff 100%);
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
                margin: 0;
                padding: 32px 0;
                color: #1c1c1e;
                display: flex;
                justify-content: center;
                min-height: 100vh;
            }

            .card {
                position: relative;
                overflow: hidden;
                background: linear-gradient(145deg, rgba(255,255,255,.76), rgba(255,255,255,.46));
                border: 1px solid rgba(255,255,255,.84);
                border-radius: 28px;
                box-shadow: 0 28px 80px rgba(48, 88, 138, .22), inset 0 1px 0 rgba(255,255,255,.9);
                backdrop-filter: blur(34px) saturate(170%);
                -webkit-backdrop-filter: blur(34px) saturate(170%);
                padding: 24px;
                width: 92%; /* 移动端占满宽度 */
                max-width: 500px; /* 桌面端限制宽度 */
                margin: auto;
                transition: transform 0.3s ease;
            }
            .card::before { content:""; position:absolute; inset:0; pointer-events:none; background:linear-gradient(120deg,rgba(255,255,255,.42),transparent 32%,rgba(255,255,255,.12) 70%,transparent); }
            .card > * { position:relative; z-index:1; }

            h2 {
                margin-top: 0;
                font-weight: 700;
                text-align: center;
                margin-bottom: 25px;
                font-size: 22px;
            }
            .badge {
                background: #007aff;
                color: white;
                padding: 3px 8px;
                border-radius: 6px;
                font-size: 12px;
                vertical-align: middle;
                margin-left: 5px;
            }

            label {
                display: block;
                font-weight: 600;
                margin-bottom: 10px;
                color: #3a3a3c;
                font-size: 15px;
            }

            /* 大触控区域优化 */
            select, input[type=file], input[type=number], input[type=text] {
                width: 100%;
                padding: 14px;
                margin-bottom: 20px;
                border: 1px solid #e5e5ea;
                border-radius: 12px;
                font-size: 16px; /* 防止iOS缩放 */
                background: rgba(255,255,255,.56);
                border-color: rgba(255,255,255,.78);
                box-shadow: inset 0 1px 0 rgba(255,255,255,.8), 0 6px 18px rgba(58,91,130,.06);
                appearance: none; /* 去除原生样式 */
                outline: none;
            }

            /* 按钮优化 */
            button {
                background-color: #0071e3;
                color: white;
                border: none;
                padding: 16px;
                font-size: 17px;
                font-weight: 600;
                border-radius: 12px;
                cursor: pointer;
                width: 100%;
                box-shadow: 0 4px 12px rgba(0,113,227,0.3);
                transition: all 0.2s;
            }
            button:active { transform: scale(0.97); opacity: 0.9; }
            button:disabled { background-color: #aeaeb2; box-shadow: none; }

            /* 进度面板优化 */
            #progressPanel { display: none; margin-top: 25px; border-top: 1px solid #f2f2f7; padding-top: 20px; }
            .progress-header { display: flex; justify-content: space-between; margin-bottom: 8px; font-size: 13px; color: #8e8e93; font-weight: 500; }
            .progress-bar-bg { background: #e5e5ea; height: 8px; border-radius: 4px; overflow: hidden; margin-bottom: 15px; }
            .progress-bar-fill { background: #34c759; height: 100%; width: 0%; transition: width 0.2s linear; }

            #logBox {
                background: #1c1c1e;
                color: #34c759;
                font-family: "Menlo", monospace;
                font-size: 11px;
                padding: 12px;
                border-radius: 10px;
                height: 180px;
                overflow-y: auto;
                line-height: 1.6;
                -webkit-overflow-scrolling: touch; /* iOS顺滑滚动 */
            }
            .log-line { margin: 0; border-bottom: 1px solid #2c2c2e; padding: 2px 0; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
            .log-line.error { color: #ff453a; }
            .options { display:grid; grid-template-columns: 1fr 1fr; gap: 0 12px; }
            .options label { font-size: 13px; margin-bottom: 4px; }
            .options input, .options select { margin-bottom: 12px; padding: 10px; }
            .full { grid-column: 1 / -1; }
            .check { display:flex; align-items:center; gap:8px; margin: 4px 0 14px; font-size:14px; }
            .titlebar { display:flex; justify-content:space-between; align-items:start; gap:12px; }
            .language, .github { width:auto; padding:8px 11px; font-size:13px; box-shadow:none; white-space:nowrap; }
            .export-row { display:flex; gap:8px; margin-bottom:18px; }
            .export-row input { margin:0; min-width:0; }
            .export-row button { width:auto; flex:0 0 auto; padding:11px 14px; font-size:14px; box-shadow:none; }
            #resultPanel { display:none; margin-top:14px; padding:14px; border-radius:14px; background:rgba(52,199,89,.1); border:1px solid rgba(52,199,89,.24); }
            #resultPanel p { margin:0 0 8px; font-size:13px; color:#28713c; word-break:break-all; }
            #resultList { margin:0 0 10px; padding-left:20px; font-size:13px; }
            #openFolderButton { width:auto; padding:9px 12px; font-size:13px; box-shadow:none; background:rgba(0,113,227,.9); }
            @media (max-width: 620px) { body { padding: max(42px, env(safe-area-inset-top)) 0 20px; } .card { border-radius:22px; padding:20px; } }
        </style>
    </head>
    <body>
        <div class="card">
            <div class="titlebar"><h2 id="appTitle">排版工厂 <span class="badge">v17.0</span></h2><div><button class="github" type="button" id="githubButton">GitHub</button> <button class="language" type="button" id="languageToggle">English</button></div></div>

            <form id="uploadForm">
                <input type="hidden" name="language" id="languageInput" value="zh">
                <label id="paperLabel">1. 纸张规格</label>
                <select id="paperSize" name="paper_size">
                    <option value="A4_STD">A4 标准 (内1.5 / 外0.5)</option>
                    <option value="A4_EQUAL">A4 等宽 (四边0.7) ✨</option>
                    <option value="B5">B5 小册子 (四边0.7)</option>
                </select>

                <label id="layoutLabel">2. 排版规则</label>
                <div class="options">
                    <label>字体大小（pt）<input type="number" name="font_size" min="6" max="72" step="0.5" value="10.5"></label>
                    <label>行间距（倍）<input type="number" name="line_spacing" min="0.5" max="5" step="0.05" value="1"></label>
                    <label>段前（pt）<input type="number" name="space_before" min="0" max="72" step="0.5" value="0"></label>
                    <label>段后（pt）<input type="number" name="space_after" min="0" max="72" step="0.5" value="0"></label>
                    <label class="full">标点转换
                        <select name="punctuation"><option value="halfwidth">转半角（默认）</option><option value="fullwidth">转全角</option><option value="preserve">保留原样</option></select>
                    </label>
                    <label class="full check"><input type="checkbox" name="remove_empty" checked> 移除没有文字或图片的空段落</label>
                    <label class="full check"><input type="checkbox" name="custom_margins"> 使用自定义边距（cm）</label>
                    <label class="full">页边距预设
                        <select name="margin_preset"><option value="paper">跟随纸张方案</option><option value="all_05">四边都是 0.5 cm</option><option value="all_07">四边都是 0.7 cm</option><option value="symmetric">对称页：内 1.5 cm，外/上下 0.7 cm</option></select>
                    </label>
                    <label>上<input type="number" name="top_margin" min="0.1" max="8" step="0.1" value="0.7"></label>
                    <label>下<input type="number" name="bottom_margin" min="0.1" max="8" step="0.1" value="0.7"></label>
                    <label>左<input type="number" name="left_margin" min="0.1" max="8" step="0.1" value="0.7"></label>
                    <label>右<input type="number" name="right_margin" min="0.1" max="8" step="0.1" value="0.7"></label>
                    <label class="full">页脚
                        <select name="footer_mode"><option value="first_line">首段文字 + 页码</option><option value="page_number">仅页码</option><option value="none">不添加页脚</option></select>
                    </label>
                    <label class="full">自定义页脚文字（留空时使用首段）<input type="text" name="footer_text" maxlength="120" placeholder="例如：课程作业"></label>
                </div>

                <label id="fileLabel">3. 选择文件 (多选)</label>
                <input type="file" id="fileInput" name="file" multiple accept=".docx" required>

                <label id="exportLabel">4. 导出位置</label>
                <div class="export-row">
                    <input type="text" id="exportDirectory" readonly value="下载文件夹">
                    <button type="button" id="chooseExportButton">选择…</button>
                </div>

                <button type="submit" id="startBtn">🚀 开始转换</button>
            </form>

            <div id="progressPanel">
                <div class="progress-header">
                    <span id="progressText">准备中...</span>
                    <span id="percentText">0%</span>
                </div>
                <div class="progress-bar-bg">
                    <div id="progressFill" class="progress-bar-fill"></div>
                </div>
                <div id="logBox"></div>
                <div id="resultPanel">
                    <p id="resultDirectory"></p>
                    <ul id="resultList"></ul>
                    <button type="button" id="openFolderButton">在文件夹中显示</button>
                </div>
            </div>
        </div>

        <script>
            const form = document.getElementById('uploadForm');
            const startBtn = document.getElementById('startBtn');
            const logBox = document.getElementById('logBox');
            const progressPanel = document.getElementById('progressPanel');
            const progressFill = document.getElementById('progressFill');
            const progressText = document.getElementById('progressText');
            const percentText = document.getElementById('percentText');
            const languageInput = document.getElementById('languageInput');
            const languageToggle = document.getElementById('languageToggle');
            const githubButton = document.getElementById('githubButton');
            const exportDirectory = document.getElementById('exportDirectory');
            const chooseExportButton = document.getElementById('chooseExportButton');
            const resultPanel = document.getElementById('resultPanel');
            const resultDirectory = document.getElementById('resultDirectory');
            const resultList = document.getElementById('resultList');
            const openFolderButton = document.getElementById('openFolderButton');
            const translations = {
                zh: { title:'排版工厂', paper:'1. 纸张规格', layout:'2. 排版规则', files:'3. 选择文件（多选）', start:'🚀 开始转换', font:'字体大小（pt）', spacing:'行间距（倍）', before:'段前（pt）', after:'段后（pt）', punctuation:'标点转换', remove:' 移除没有文字或图片的空段落', margins:' 使用自定义边距（cm）', marginPreset:'页边距预设', top:'上', bottom:'下', left:'左', right:'右', footer:'页脚', customFooter:'自定义页脚文字（留空时使用首段）', half:'转半角（默认）', full:'转全角', preserve:'保留原样', first:'首段文字 + 页码', page:'仅页码', none:'不添加页脚', paperPreset:'跟随纸张方案', all05:'四边都是 0.5 cm', all07:'四边都是 0.7 cm', symmetric:'对称页：内 1.5 cm，外/上下 0.7 cm', placeholder:'例如：课程作业', selectFiles:'请先选择文件！', processing:'处理中...', connect:'连接服务器...', connection:'连接失败', downloading:'✅ 完成，正在下载...', complete:'转换完成', retry:'重试', ready:'准备中...', toggle:'English', github:'GitHub' },
                en: { title:'AutoWord Formatter', paper:'1. Paper size', layout:'2. Formatting rules', files:'3. Select files (multiple)', start:'🚀 Start formatting', font:'Font size (pt)', spacing:'Line spacing', before:'Space before (pt)', after:'Space after (pt)', punctuation:'Punctuation', remove:' Remove empty paragraphs without text or images', margins:' Use custom margins (cm)', marginPreset:'Margin preset', top:'Top', bottom:'Bottom', left:'Left', right:'Right', footer:'Footer', customFooter:'Custom footer (uses first paragraph when blank)', half:'Convert to half-width (default)', full:'Convert to full-width', preserve:'Keep unchanged', first:'First paragraph + page number', page:'Page number only', none:'No footer', paperPreset:'Use paper preset', all05:'0.5 cm on all sides', all07:'0.7 cm on all sides', symmetric:'Mirrored: 1.5 cm inner, 0.7 cm outer/top/bottom', placeholder:'For example: Course assignment', selectFiles:'Select at least one file first.', processing:'Processing...', connect:'Connecting to local service...', connection:'Connection failed', downloading:'✅ Done. Downloading...', complete:'Formatting complete', retry:'Retry', ready:'Preparing...', toggle:'中文', github:'GitHub' }
            };
            Object.assign(translations.zh, { export:'4. 导出位置', choose:'选择…', reveal:'在文件夹中显示', saved:'已保存到：', downloading:'✅ 完成，正在保存文件...' });
            Object.assign(translations.en, { export:'4. Export location', choose:'Choose…', reveal:'Show in folder', saved:'Saved to: ', downloading:'✅ Done. Saving files...' });
            let language = localStorage.getItem('autoword-language') || 'zh';
            const labelText = (name, value) => document.querySelector(`[name="${name}"]`).parentElement.childNodes[0].nodeValue = value;
            function setLanguage(next) {
                language = next; localStorage.setItem('autoword-language', language); languageInput.value = language;
                const t = translations[language]; document.documentElement.lang = language === 'en' ? 'en' : 'zh-CN'; document.title = t.title;
                document.getElementById('appTitle').childNodes[0].nodeValue = t.title + ' ';
                document.getElementById('paperLabel').textContent = t.paper; document.getElementById('layoutLabel').textContent = t.layout; document.getElementById('fileLabel').textContent = t.files;
                document.getElementById('exportLabel').textContent = t.export; chooseExportButton.textContent = t.choose; openFolderButton.textContent = t.reveal;
                labelText('font_size', t.font); labelText('line_spacing', t.spacing); labelText('space_before', t.before); labelText('space_after', t.after); labelText('top_margin', t.top); labelText('bottom_margin', t.bottom); labelText('left_margin', t.left); labelText('right_margin', t.right); labelText('footer_text', t.customFooter);
                document.querySelector('[name="punctuation"]').parentElement.childNodes[0].nodeValue = t.punctuation; document.querySelector('[name="footer_mode"]').parentElement.childNodes[0].nodeValue = t.footer;
                document.querySelector('[name="margin_preset"]').parentElement.childNodes[0].nodeValue = t.marginPreset;
                document.querySelector('[name="remove_empty"]').parentElement.childNodes[1].nodeValue = t.remove; document.querySelector('[name="custom_margins"]').parentElement.childNodes[1].nodeValue = t.margins;
                const punct = document.querySelector('[name="punctuation"]').options; [punct[0].text, punct[1].text, punct[2].text] = [t.half, t.full, t.preserve];
                const footer = document.querySelector('[name="footer_mode"]').options; [footer[0].text, footer[1].text, footer[2].text] = [t.first, t.page, t.none];
                const margins = document.querySelector('[name="margin_preset"]').options; [margins[0].text, margins[1].text, margins[2].text, margins[3].text] = [t.paperPreset, t.all05, t.all07, t.symmetric];
                document.querySelector('[name="footer_text"]').placeholder = t.placeholder; if (!startBtn.disabled) startBtn.textContent = t.start; progressText.textContent = t.ready; languageToggle.textContent = t.toggle; githubButton.textContent = t.github;
            }
            languageToggle.addEventListener('click', () => setLanguage(language === 'zh' ? 'en' : 'zh'));
            githubButton.addEventListener('click', () => window.open('https://github.com/xiaozhangwangxue/autoword', '_blank', 'noopener'));
            window.addEventListener('pywebviewready', async () => { exportDirectory.value = await window.pywebview.api.get_default_export_directory(); });
            chooseExportButton.addEventListener('click', async () => {
                if (window.pywebview?.api) exportDirectory.value = await window.pywebview.api.choose_export_directory();
            });
            openFolderButton.addEventListener('click', async () => {
                if (window.pywebview?.api) await window.pywebview.api.open_export_directory(exportDirectory.value);
            });
            document.querySelector('[name="margin_preset"]').addEventListener('change', event => {
                const values = { all_05:['0.5','0.5','0.5','0.5'], all_07:['0.7','0.7','0.7','0.7'], symmetric:['0.7','0.7','1.5','0.7'] }[event.target.value];
                if (!values) return;
                ['top_margin','bottom_margin','left_margin','right_margin'].forEach((name, index) => document.querySelector(`[name="${name}"]`).value = values[index]);
                document.querySelector('[name="custom_margins"]').checked = false;
            });
            setLanguage(language);

            form.addEventListener('submit', async function(e) {
                e.preventDefault();

                const files = document.getElementById('fileInput').files;
                const t = translations[language];
                if (files.length === 0) { alert(t.selectFiles); return; }

                startBtn.disabled = true;
                startBtn.textContent = t.processing;
                progressPanel.style.display = 'block';
                resultPanel.style.display = 'none';
                logBox.innerHTML = '<div class="log-line">> ' + t.connect + '</div>';

                const formData = new FormData(form);
                let completed = false;

                try {
                    const response = await fetch('/process_stream', { method: 'POST', body: formData });
                    if (!response.ok) throw new Error(t.connection);

                    const reader = response.body.getReader();
                    const decoder = new TextDecoder();
                    let buffer = '';

                    while (true) {
                        const { done, value } = await reader.read();
                        if (done) break;

                        buffer += decoder.decode(value, { stream: true });
                        const lines = buffer.split('\\n');
                        buffer = lines.pop();

                        for (const line of lines) {
                            if (!line.trim()) continue;
                            try {
                                const data = JSON.parse(line);

                                if (data.status === 'log') {
                                    const div = document.createElement('div');
                                    div.className = 'log-line';
                                    div.textContent = "> " + data.msg;
                                    logBox.appendChild(div);
                                    logBox.scrollTop = logBox.scrollHeight;
                                }
                                else if (data.status === 'progress') {
                                    const pct = Math.round(data.val * 100);
                                    progressFill.style.width = pct + '%';
                                    percentText.textContent = pct + '%';
                                    progressText.textContent = data.msg;
                                }
                                else if (data.status === 'done') {
                                    completed = true;
                                    const div = document.createElement('div');
                                    div.className = 'log-line';
                                    div.style.color = '#34c759';
                                    div.textContent = "> " + t.downloading;
                                    logBox.appendChild(div);

                                    if (window.pywebview?.api) {
                                        const result = await window.pywebview.api.export_job(data.job_id, exportDirectory.value);
                                        if (!result.ok) throw new Error(result.error || t.connection);
                                        exportDirectory.value = result.directory;
                                        resultDirectory.textContent = t.saved + result.directory;
                                        resultList.innerHTML = '';
                                        result.files.forEach(file => { const li = document.createElement('li'); li.textContent = file.name; resultList.appendChild(li); });
                                        resultPanel.style.display = 'block';
                                    } else {
                                        window.location.href = "/download/" + data.job_id;
                                    }

                                    startBtn.textContent = t.complete;
                                    startBtn.style.backgroundColor = "#34c759";
                                    document.getElementById('fileInput').value = '';
                                    window.setTimeout(() => {
                                        startBtn.disabled = false;
                                        startBtn.textContent = translations[language].start;
                                        startBtn.style.backgroundColor = '';
                                    }, 900);
                                }
                                else if (data.status === 'error') {
                                    const div = document.createElement('div');
                                    div.className = 'log-line error';
                                    div.textContent = data.msg;
                                    logBox.appendChild(div);
                                }
                            } catch (err) { }
                        }
                    }
                    if (!completed) {
                        startBtn.disabled = false;
                        startBtn.textContent = t.retry;
                        startBtn.style.backgroundColor = '';
                    }
                } catch (error) {
                    alert("错误: " + error.message);
                    startBtn.disabled = false;
                    startBtn.textContent = t.retry;
                    startBtn.style.backgroundColor = '';
                }
            });
        </script>
    </body>
    </html>
    '''

@app.route('/process_stream', methods=['POST'])
def process_stream():
    uploaded_files = request.files.getlist('file')
    settings = layout_settings(request.form)

    valid_files = [f for f in uploaded_files if f.filename and f.filename.lower().endswith('.docx')]

    if not valid_files:
        return Response(json.dumps({"status": "error", "msg": "❌ 无有效docx文件"}) + "\n", mimetype='application/json')

    job_id = str(uuid.uuid4())
    job_dir = os.path.join(UPLOAD_FOLDER, job_id)
    os.makedirs(job_dir, exist_ok=True)

    # FileStorage streams are closed by Flask after the request returns. Save
    # every upload before creating the streaming response so packaged desktop
    # apps do not fail midway with a closed-file 500 error.
    local_paths = []
    for uploaded_file in valid_files:
        filename = os.path.basename(uploaded_file.filename.replace('\\', '/'))
        if not filename or filename in {'.', '..'} or '\x00' in filename:
            continue
        path = os.path.join(job_dir, f"{len(local_paths)}_{filename}")
        uploaded_file.save(path)
        local_paths.append((path, filename))

    if not local_paths:
        return Response(json.dumps({"status": "error", "msg": "❌ 无有效docx文件"}) + "\n", mimetype='application/json')

    def generate():
        total = len(local_paths)
        processed_files = []

        for i, (path, filename) in enumerate(local_paths):
            yield json.dumps({"status": "progress", "val": i / total, "msg": f"处理 {i+1}/{total}"}) + "\n"

            gen = process_single_file_task(path, settings, job_id, filename)

            file_data = None
            file_name = None
            for item in gen:
                if isinstance(item, str): yield item
                elif isinstance(item, dict):
                    file_name = item['filename']
                    file_data = item['data']

            if file_data:
                processed_files.append((file_name, file_data))
                yield json.dumps({"status": "progress", "val": (i + 0.9) / total, "msg": f"完成: {filename}"}) + "\n"

        yield json.dumps({"status": "log", "msg": "正在打包..."}) + "\n"

        if not processed_files:
            yield json.dumps({
                "status": "error",
                "msg": localized(settings, "❌ 所有文件转换失败，请查看上方错误信息", "❌ All files failed. See the error above."),
            }) + "\n"
            return

        zip_path = os.path.join(job_dir, f"result_{job_id}.zip")
        output_dir = os.path.join(job_dir, "output")
        os.makedirs(output_dir, exist_ok=True)
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for fname, fdata in processed_files:
                zf.writestr(fname, fdata)
                with open(os.path.join(output_dir, fname), 'wb') as output_file:
                    output_file.write(fdata)

        yield json.dumps({"status": "progress", "val": 1.0, "msg": "完成"}) + "\n"
        yield json.dumps({"status": "done", "job_id": job_id}) + "\n"

    return Response(stream_with_context(generate()), mimetype='application/x-ndjson')

@app.route('/download/<job_id>')
def download_result(job_id):
    job_dir = os.path.join(UPLOAD_FOLDER, job_id)
    zip_path = os.path.join(job_dir, f"result_{job_id}.zip")
    if os.path.exists(zip_path):
        return send_file(zip_path, as_attachment=True, download_name=f'排版成品.zip')
    else:
        return "Expired", 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
